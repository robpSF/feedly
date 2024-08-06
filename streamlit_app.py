# GET NEW TOKENS HERE
# https://feedly.com/v3/auth/dev#


import SL_GET_SM
import robslib
import requests, json
import pandas as pd
import docx
from docx import Document
import html2text as ht
import re
from datetime import datetime

from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_BREAK


path ="C:/Users/rober_vsah/Downloads/"
outputfile = "test2.docx"
document = Document()


board = "https://feedly.com/i/board/content/user/e32c40ed-b126-4ca4-9609-8cf10fcfa9b9/tag/7dbae0c0-1f31-4c93-80aa-fe06c0de382c"
consumer_key = "e32c40ed-b126-4ca4-9609-8cf10fcfa9b9"

def get_feedly_articles(consumer_key):
    url = "https://cloud.feedly.com/v3/boards"
    headers = {"Authorization":"Bearer "+consumer_key}
    response = requests.get(url, headers=headers)
    data = json.loads(response.content)

    return data

import json
import requests
import streamlit as st

# Feedly
access_token=st.secrets["feedly"]["access_token"]

feedaccess = access_token
myfeedid = "e32c40ed-b126-4ca4-9609-8cf10fcfa9b9"
feedcount = "20"

def get_boards():
    url = "https://cloud.feedly.com/v3/boards?withEnterprise=True"
    headers = {'Authorization': 'OAuth ' + feedaccess}
    response = requests.get(url, headers=headers)
    data = json.loads(response.content)
    st.write("board response",response)
    #st.write(data)
    df = pd.DataFrame.from_dict(data)
    st.write(df)
    boards = df["label"].tolist()
    return boards,df

def refresh_feedly_token():
    url = "https://cloud.feedly.com/v3/auth/token"
    headers = {'Authorization': 'OAuth ' + feedaccess}
    params = {
        "refresh_token":"A24pqm5fvqx9o0fQh_dzo8bYdTAfcjzx3Dqh1e0dmAm_DptAmq4ouVkd88SqbPKD6S3z2t6c-rha5NU8eU62VMfnx_HJoHtX6aoIf-hNC2g__Kj1-4s-s_mLWckc3ZpLIz-eX5IZqgekWmSSRw2wOdUFyIy4I7Fx6ReQmgs3m7Rtxp4hfgBsfoklRLtjuBGmoqSxsrQ4Kptjrs1n2xVQIZ7t58BMPzyWd4li8fPfwAgJfH3dqzWW2X0XgAQGu6Gr:feedlydev",
            "client_id": "e32c40ed-b126-4ca4-9609-8cf10fcfa9b9",
    "client_secret": access_token,
    "grant_type":"refresh_token"
    }
    response = requests.post(url,params=params, headers=headers)
    data = json.loads(response.content)
    print("Feedly refresh",data)
    return data


def get_articles(board_id):
    st.write(board_id)
    url2 = "https://cloud.feedly.com/v3/streams/contents"
    params = dict(streamId=board_id)
    headers = {'Authorization': 'OAuth ' + feedaccess}
    response = requests.get(url2, params=params, headers=headers)
    st.write("get articles response",response)
    data2 = json.loads(response.content)
    content_list= data2["items"]
    st.write(content_list)
    return content_list

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value

    part = paragraph.part
    #print(url)
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    #print(r_id)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    #print(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def replace_with_html(master_content,start):
    #assumes that the first and last character on the highlighted text is either _ or * for italic and bold
    #start is the character index where the highlighted text begins
    highlighted_text= master_content[start:len(master_content)]
    st.header("Here!")
    st.write(highlighted_text)
    quit(500)
    if highlighted_text[0]== "_": #this it's italic
        master_content=master_content[:start-1]+"<i>" #chop off the highlight and add <i>
        master_content=master_content+highlighted_text[1:-1] #all but the last characters
        master_content=master_content+"</i>" #all done
        st.write(master_content)
        quit(500)
    return master_content

def print_articles(write_to_doc, content_list, document):

    # gather just the bits we need

    df = pd.DataFrame(columns = ("Serial", \
                                 "Number", \
                        "Time",\
                        "From", \
                        "Faction",\
                        "To", \
                        "Team", \
                        "Method", \
                        "Subject", \
                        "Message", \
                        "Timestamp",\
                        "Expected Action",\
                        "ImageURL"))

    for item in content_list:
        persona = item["origin"]["title"]
        title = item["title"]


        #some tweaks to fix persona issues
        if persona.find("Google News") >= 0:
            persona = "Google News"

        if persona.find("Sputnik News") >= 0:
            persona = "Sputnik"

        if persona.find("Mail Online") >= 0:
            persona = "Mail Online"

        if persona.find("SvD") >= 0:
            persona = "SvD"

        if persona.find("NRC") >= 0:
            persona = "NRC"

        if persona.find("lrytas") >= 0:
            persona = "Lrytas.it"

        if persona.find("berlingske") >= 0:
            persona = "Berlingske"

        article_url = item["alternate"][0]["href"]

        #fix the timestamp
        #try:
        #discovered that I needed to divide by 1000
        #https://stackoverflow.com/questions/37494983/python-fromtimestamp-oserror
        unix = item["published"]/1000
        #st.sidebar.write(unix)
        publication_date = datetime.fromtimestamp(unix).strftime('%Y-%m-%dT%H:%M')
        #st.sidebar.write(publication_date)
        #except:
        #    st.sidebar.write("Missing publication date")
        #    publication_date = ""

        #get the content
        try:
            content = ht.html2text(item["fullContent"])
        except:
            # if full content not available, get the summary
            try:
                content = ht.html2text(item["summary"]["content"])
            except:
                content = ""

        #strip URLs from text
        content = re.sub(r'^https?:\/\/.*[\r\n]*', '', content, flags=re.MULTILINE)
        #try again
        #content = re.sub(r'\w+:\/{2}[\d\w-]+(\.[\d\w-]+)*(?:(?:\/[^\s/]*))*', '', content)
        # remove content between () and []
        content = re.sub(r"\([^()]*\)", "", content)

        # now remove any remaining [ and ]
        content = content.replace("[", " ")
        content = content.replace("]", " ")


        # now remove the end of lines
        length = len(content)
        #print(repr(content))
        new_content = ""
        i=0
        start = 0
        while i<length:
            try:
                if (content[i] == '\n'):
                    if (content[i+1] == "\n"): #do nothing because end of paragraph
                        #st.write("im here!")
                        #st.write(new_content)
                        new_content = new_content + "\n\n"
                        i = i+2
                    else: #replace
                        new_content = new_content+" "
                        i=i+1
                elif content[i]==">":
                    new_content = new_content + ""
                    i = i + 1
                elif content[i:i+1] == "**":
                    st.sidebar("_")
                    if start>0:
                        #add the _ to be replaced
                        new_content = new_content + content[i]
                        i = i + 1
                        new_content = replace_with_html(new_content,start)
                    else:
                        start = i
                        new_content = new_content + content[i]
                        i = i + 1
                else:
                    new_content = new_content + content[i]
                    i = i + 1

            except:
                break


        #st.header("Original")
        #st.write(content)
        #st.header("New")
        content = new_content
        #st.write(content)
        #print("new")
        #print(repr(content))

        no_image= True
        #get an image
        try:
            image = item["visual"]["url"]
        except:
            try:
                image = item["enclosure"][0]["href"]
                no_image = False
            except:
                no_image=True
                image = ""

        # This is for the Conducttr import
        df = df.append({"Serial": "News ",\
                        "Number":"",\
                        "Time":"09:00:00",\
                        "From":persona, \
                        "Faction":"AI Media",\
                        "To": "All", \
                        "Team": "", \
                        "Method": "Websites / News", \
                        "Subject": title, \
                        "Message": content, \
                        "Timestamp": publication_date,\
                        "Expected Action" :"",\
                        "ImageURL": image}, ignore_index=True)


        nuggets = []
        comments = []
        try:
            for leo in item["leoSummary"]["sentences"]:
                nuggets.append(leo["text"])
        except:
            do_nothing = True

        try:
            for comment in item["annotations"]:
                comments.append(comment["comment"])
        except:
            do_nothing = True

        # now write it out!
        st.header(persona)
        st.subheader(title)
        if not no_image: st.image(image,width=250)
        if comments != []:
            # st.markdown("**Comment**")
            for comment in comments:
                st.markdown("**" + comment + "**")
        for nugget in nuggets:
            st.write(nugget)
        st.write(content)
        st.write(article_url)

        if write_to_doc:
            document.add_paragraph(title)
            if comments != []:
                # st.markdown("**Comment**")
                for comment in comments:
                    document.add_paragraph(comment)
            for nugget in nuggets:
                document.add_paragraph(nugget)
            p = document.add_paragraph()
            add_hyperlink(p, article_url, article_url)

    st.table(df)
    return df


boards,df = get_boards()

# pick a board and make it look nice
this_board = st.sidebar.selectbox(label="Pick a board",options=boards)

board_id = df.loc[df["label"]==this_board,"id"].iloc[0]
cover_image = ""# df.loc[df["label"]==this_board,"cover"].iloc[0]
st.header(this_board)
if cover_image != None:
    try:
        st.image(cover_image,width=600)
    except:
        do_nothing = True
#st.write(board_id)

articles = get_articles(board_id)
df = print_articles(False, articles, document)

if st.sidebar.button(label="Save file"):
    print_articles(True, articles, document)
    document.save(path + outputfile)
    st.sidebar.text("Saved document "+ outputfile+ " in "+ path)
    print("Saved document ", outputfile, " in ", path)
    df.to_excel(path+"news_scrape.xlsx",encoding="utf-8", index=False)
    df.to_csv(path + "news_scrape.csv", encoding="utf-8",index=False)
    st.sidebar.text("News saved to news_scrape.xlsx in " + path)
    st.header("Done! Files saved")
